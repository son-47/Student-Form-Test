from odoo import http
from odoo.http import request,Response, content_disposition
import json,io
import os,glob 
import re
import base64
import shutil, uuid
import mimetypes 
from datetime import datetime, date
import logging
from docx import Document
import pandas as pd
from werkzeug.utils import secure_filename
_logger = logging.getLogger(__name__)
from ..validator.student_validator import StudentValidator
# import cloudinary
# import cloudinary.uploader,cloudinary.api
# cloudinary.config(
#     cloud_name= 'dg4htxiop',
#     api_key= '539252181871726',
#     api_secret= 'z14rDY3Yvnwh7DIgDP1Rzny69qQ'
# )




MODULE_DIR = os.path.dirname(os.path.abspath(__file__))  # /addons/my_api_module/controllers/
BASE_DIR = os.path.abspath(os.path.join(MODULE_DIR, '..'))  # dẫn lên thư mục cha /addons/my_api_module/ 
DOCS_DIR = os.path.join(BASE_DIR, 'static', 'uploads', 'docs')  # Đường dẫn tới static/uploads/docs
IMAGES_DIR = os.path.join(BASE_DIR, 'static', 'uploads', 'images')  # Đường dẫn tới static/uploads/images 
BASE_URL = "http://172.27.100.251:8069"


ALL_HOBBIES = ['Astronomy', 'Badminton', 'Basketball', 'Coding','Collecting', 'Cooking', 'Crafting',
               'Cycling', 'Dancing', 'Drawing', 'Fishing', 'Fitness', 'Football', 'Gaming', 'Gardening',
               'Hiking', 'Music', 'Other', 'Photography', 'Reading', 'Running', 'Swimming',
               'Table Tennis', 'Tennis', 'Traveling', 'Volunteering', 'Volleyball', 'Writing',
               'Yoga']
## type : lấy dữ liệu đầu vào từ đâu
## auth : thiết lập quyền truy cập
########: # public: không cần đăng nhập
    #         # user: cần đăng nhập
    #         # 'bearer': (API bảo mật nâng cao)
    #         # 'none' : Không liên quan gì đến user hay database. 
    #           Dùng trong các tình huống hệ thống chưa khởi tạo database  
    
## csrf : Cross-Site Request Forgery, bảo vệ chống lại các cuộc tấn công giả mạo yêu cầu từ trang web khác
# khi tắt thì Cho phép gửi request từ client khác (Postman, mobile app)

 # Thêm nếu chưa có

class StudentController(http.Controller):
    @http.route('/images/<filename>', type='http', auth='public', methods=['GET'])
    def serve_uploaded_file(self, filename):
        """Phục vụ file từ thư mục IMAGES_DIR"""
        filepath = os.path.join(IMAGES_DIR, filename)
        if not os.path.isfile(filepath):
            return request.make_response(
                json.dumps({
                    "code": "404",
                    "status": "error",
                    "message": "File không tồn tại",
                    "data": {}
                }),
                headers=[('Content-Type', 'application/json')],
                status=404
            )
        try:
            with open(filepath, 'rb') as f:
                file_content = f.read()
            # Xác định MIME type dựa trên phần mở rộng
            mime_type = 'image/jpeg' if filename.endswith('.jpg') or filename.endswith('.jpeg') else 'image/png'
            return request.make_response(
                file_content,
                headers=[
                    ('Content-Type', mime_type),
                    ('Content-Disposition', f'inline; filename="{filename}"')
                ]
            )
        except Exception as e:
            _logger.error(f"Lỗi khi phục vụ file {filename}: {str(e)}")
            return request.make_response(
                json.dumps({
                    "code": "500",
                    "status": "error",
                    "message": f"Lỗi khi đọc file: {str(e)}",
                    "data": {}
                }),
                headers=[('Content-Type', 'application/json')],
                status=500
            )

class StudentController(http.Controller):

    def encode_hobbies_binary_string_to_bitmask(self, hobbies_input):
        """
        Convert hobbies input to a bitmask (integer).
        - If input is a binary string (e.g., '0,1,1,0,...'), validate and convert to bitmask.
        - If input is a comma-separated list of hobbies (e.g., 'Badminton,Basketball'), convert to binary string then to bitmask.
        """
     
        bits = list(map(int, hobbies_input.strip().split(',')))
        bitmask = 0 
        for i, value in enumerate(bits):
            if value == 1:
                bitmask |= (1 << i) # giữ nguyên hoặc bằng 2^i
        return bitmask

       
    
    # ##decode hobbies về chuỗi 0,1,...
    def decode_bitmask_to_hobbies_string(self, bitmask):
        bits = [(bitmask >> i) & 1 for i in range(len(ALL_HOBBIES))]
        return ','.join(str(b) for b in bits)

        
    ###1: LẤY DANH SÁCH TẤT CẢ SINH VIÊN
    @http.route('/tra_student', type='http', auth='public', methods=['GET'], csrf=False)
    def get_all(self, **kw):
        try:
            columnlist = kw.get('columnlist')
            output_fields = [
            'id', 'code', 'fullname', 'dob', 'sex', 'homecity', 'address',
            'hobbies', 'hair_color', 'email', 'facebook', 'class_id',
            'username', 'password', 'description', 'attachment','attachment_url'
        ]
            fields = columnlist.split(',') if columnlist else output_fields

            students = request.env['my_api_module.student'].search([])

            result = []
            for student in students:
                record = {}
                for field in fields:
                    if hasattr(student, field):
                        value = getattr(student, field)
                        if isinstance(value, date):
                            record[field] = value.strftime('%Y-%m-%d')
                        elif field == 'class_id': 
                            record[field] = value.id if value else None
                        elif field == 'hobbies':
                            record[field] = self.decode_bitmask_to_hobbies_string(value)
                        # elif field == "attachment":
                        #     record[field] = ""
                        else:
                            record[field] = value
                result.append(record)

            return request.make_response(
                json.dumps({
                    "code": "200",
                    "status": "success",
                    "message": "Lấy danh sách sinh viên thành công",
                    "data": result
                }),
                headers=[('Content-Type', 'application/json')]
            )

        except Exception as e:
            return request.make_response(
                json.dumps({
                    "code": "B600",
                    "status": "error",
                    "message": f"Lỗi không xác định: {str(e)}",
                    "data": []
                }),
                headers=[('Content-Type', 'application/json')]
            )

##2: LẤY DANH SÁCH SINH VIÊN THEO TIÊU CHÍ
    @http.route('/tra_student/page/<int:init>', type='http', auth='public', methods=['GET'], csrf=False)
    def get_by_page(self, init, **kw):
        try:
            # Validate page
            if init < 1:
                raise ValueError('C601')  # lỗi định dạng trang

            # Parse parameters
            try:
                size = int(kw.get('size', 10))
                search = kw.get('search', '').strip()
                columnlist = kw.get('columnlist')
                order_param = kw.get('order', '')
                toplist = kw.get('toplist', '')
            except:
                raise ValueError('C602')  # lỗi định dạng 

            # Validate columns
            all_fields = [
                'id', 'code', 'fullname', 'dob', 'sex', 'homecity', 'address',
                'hobbies', 'hair_color', 'email', 'facebook', 'class_id',
                'username', 'password', 'description', 'attachment','attachment_url'
            ]
            fields_list = columnlist.split(',') if columnlist else all_fields
            if not all(field in all_fields for field in fields_list):
                raise ValueError('C607')  # lỗi columnlist không hợp lệ

            # Parse toplist
            toplist_ids = []
            if toplist:
                toplist_ids = [int(x) for x in toplist.split(',') if x.strip().isdigit()]

            # Parse order: vd "co:1-na:0"
            order_list = []
            for part in order_param.split('-'):
                if ':' in part:
                    key, direction = part.split(':')
                    mapping = {'co': 'code', 'na': 'fullname'}
                    if key in mapping:
                        field = mapping[key]
                        order_list.append(f"{field} {'desc' if direction == '1' else 'asc'}")
            order_str = ', '.join(order_list) if order_list else 'id asc'

            # Xây domain tìm kiếm
            domain = []
            # ilike không phân biệt chữ hoa chữ thường
            if search:
                domain += ['|', '|', '|', '|',
                           ('code', 'ilike', search),
                           ('fullname', 'ilike', search),
                           ('email', 'ilike', search),
                           ('facebook', 'ilike', search),
                           ('address', 'ilike', search)]
            domain_for_students = domain + [('id', 'not in', toplist_ids)] if toplist_ids else domain    
            Student = request.env['my_api_module.student']

            # Tổng số bản ghi thỏa mãn domain
            total_items = Student.search_count(domain)

            # Tổng số trang
            total_pages = (total_items + size - 1) // size
            if init > total_pages and total_pages != 0:
                raise ValueError('C601') 
            # Bản ghi chính offset : bỏ qua (init - 1) * size bản ghi đầu tiên,limit lấy tối đa size bản ghi

            students = Student.search(domain_for_students, order=order_str, offset=(init - 1) * size, limit=size)


            # Bản ghi toplist
            toplist_records = Student.browse(toplist_ids) if toplist_ids else []

            # Kết hợp dữ liệu
            all_records = list(toplist_records) + list(students)

            # Xử lý dữ liệu trả về
            def serialize(record):
                result = {}
                for field in fields_list:
                    value = getattr(record, field, '')
                    # if field == 'dob':
                    #     _logger.info(f"ngày sinh dạng: {type(value)} ")
                    if isinstance(value,  date):
                        result[field] = value.strftime('%Y-%m-%d')
                    elif hasattr(value, 'id') and hasattr(value, 'name'):
                        result[field] = value.id if value else None
                    elif field == 'attachment':
                        result[field] = ""
                    elif field == 'hobbies':
                        result[field] = self.decode_bitmask_to_hobbies_string(value)
                    else:
                        result[field] = value
                return result
    
            data = {
                'records': [serialize(s) for s in all_records],
                'pageinfo': {
                    'count': total_items,
                    'current': init,
                    'total_pages': total_pages,
                    'size': size
                }
            }

            return request.make_response(
                json.dumps({
                    "code": "200",
                    "status": "success",
                    "message": "Lấy danh sách sinh viên thành công",
                    "data": data
                }),
                headers=[('Content-Type', 'application/json')]
            )

        except ValueError as ve:
            code = str(ve)
            messages = {
                'C601': "Lỗi định dạng của số trang",
                'C602': "Lỗi định dạng của Sinh viên trang",
                'C607': "Lỗi danh sách columnList"
            }
            return request.make_response(
                json.dumps({
                    "code": code,
                    "status": "error",
                    "message": messages.get(code, "Tham số không hợp lệ"),
                    "data": []
                }),
                headers=[('Content-Type', 'application/json')]
            )

        except Exception as e:
            return request.make_response(
                json.dumps({
                    "code": "C600",
                    "status": "error",
                    "message": f"Lỗi không xác định: {str(e)}",
                    "data": []
                }),
                headers=[('Content-Type', 'application/json')]
            )



##3: THÊM MỚI MỘT SINH VIÊN
### Upload lên filesystem
    @http.route('/tra_student', type='http', auth='public', methods=['POST'], csrf=False)
    def store(self, **kw):
        try:
            # --- Xử lý dữ liệu đầu vào
            if request.httprequest.content_type == 'application/json':
                data = json.loads(request.httprequest.data)
            else:
                data = kw
            _logger.info("[STORE] Dữ liệu sau khi phân tích: %s", data)
            # --- Trường bắt buộc
            required_fields = ['code', 'fullname','description', 'dob', 'sex', 'email', 'username', 'password','hobbies','hair_color','homecity','address', 'facebook','address','attachment','attachment_url']
            # missing_fields = [f for f in required_fields if not data.get(f)]
            # if missing_fields:
            #     raise ValueError("E603")
            validator = StudentValidator(kw, request.env["my_api_module.student"])
            errors = validator.validate_create_data()
            if errors:
                return request.make_response(
                    json.dumps({
                        "code": "E603",
                        "status": "error",
                        "message": "Lỗi kiểm tra dữ liệu: " + " ".join(errors),
                        "data": {}
                    }),
                    headers=[('Content-Type', 'application/json')]
                )
            # --- Chuẩn hóa ngày sinh
            try:
                dob_str = data.get('dob', '').strip()
                data['dob'] = datetime.strptime(dob_str, "%Y-%m-%d").date()
            except Exception as e:
                raise ValueError("Lỗi chuẩn hóa ngày sinh")

            #Kiểm tra hobbies
            if 'hobbies' in data:
                try:
                    data['hobbies'] = self.encode_hobbies_binary_string_to_bitmask(data['hobbies'])
                    _logger.info("[STORE] Hobbies đã được mã hóa: %s", data['hobbies'])
                except Exception as e:
                    raise ValueError("Lỗi mã hóa hobbies")            
            # --- Kiểm tra trùng mã sinh viên
            existing = request.env['my_api_module.student'].search([('code', '=', data['code'])], limit=1)
            if existing:
                raise ValueError("lỗi trùng mã sinh viên")

            # --- Encode hobbies
            
            _logger.info("In ra toàn bộ data: %s", data)
            
            # Tạo record mới, loại bỏ file attachment nếu có 
            data_copy = data.copy()
            data_copy.pop('fattachment', None)
            for field in required_fields:
                if field not in data_copy:
                    data_copy[field] = "1" if field == "sex" else ""
            try:
                student = request.env['my_api_module.student'].sudo().create(data_copy)
            except Exception as e:
                raise ValueError(f"Lỗi khi tạo sinh viên: {str(e)}")
            
            #Xử lý file upload
             
            file = request.httprequest.files.get("fattachment")
            _logger.info("[STORE] Danh sách file upload: %s", list(request.httprequest.files.keys()))
            attachment_url = ""  # Khởi tạo mặc định
            if file and hasattr(file, 'filename') and file.filename:
                _logger.info("[STORE] Nhận được file: %s", file.filename)
                # Tạo tên file an toàn và duy nhất
                timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
                filename = f"student_{timestamp}_{student.id}_{secure_filename(file.filename)}"
                filepath = os.path.join(IMAGES_DIR, filename)
                try:
                    os.makedirs(IMAGES_DIR, exist_ok=True)
                    with open(filepath, 'wb') as f:
                        file.save(f)
                        # Mới tạo test thử 
                    attachment = f"images/{filename}"
                    attachment_url = f"{BASE_URL}/images/{filename}"
                    _logger.info("[STORE] File uploaded to: %s", attachment_url)
                except Exception as e:
                    student.unlink()  # Xóa sinh viên nếu có lỗi
                    raise ValueError(f"Lỗi khi ghi file: {str(e)}")
            else:
                _logger.error(f"File không hợp lệ hoặc không tồn tại: type = {type(file)}, giá trị = {file}")
                # Không raise lỗi, vẫn tạo sinh viên bình thường nếu không có file

            # --- Tạo sinh viên ---
            try:
                student.write({'attachment_url': attachment_url, 'attachment': attachment})
                _logger.info("[STORE] Sinh viên đã được tạo: %s", student.id)
            except Exception as e:
                student.unlink()  # Xóa sinh viên nếu có lỗi
                raise ValueError(f"Lỗi khi tạo sinh viên: {str(e)}")
            return request.make_response(
                json.dumps({
                    "code": 200,
                    "status": "success",
                    "message": "Tạo mới sinh viên thành công",
                    "data": {
                        "id": student.id
                    }
                }),
                headers=[('Content-Type', 'application/json')]
            )

        except ValueError as ve:
            return request.make_response(
                json.dumps({
                    "code": "E603",
                    "status": "error",
                    "message": f"Lỗi kiểm tra dữ liệu: {str(ve)}",
                    "data": {}
                }),
                headers=[('Content-Type', 'application/json')]
            )

        except Exception as e:
            return request.make_response(
                json.dumps({
                    "code": "600",
                    "status": "error",
                    "message": f"Lỗi không xác định: {str(e)}",
                    "data": {}
                }),
                headers=[('Content-Type', 'application/json')]
            )



### Store gốc
    # @http.route('/tra_student', type='http', auth='public', methods=['POST'], csrf=False)
    # def store(self, **kw):
    #     try:
    #         # --- Xử lý dữ liệu đầu vào từ cả Postman (form-data) và Mobile App (JSON)
    #         if request.httprequest.content_type == 'application/json':
    #             data = json.loads(request.httprequest.data)
    #         else:
    #             data = kw

    #         # --- Danh sách trường bắt buộc
    #         required_fields = ['code', 'fullname', 'dob', 'sex', 'email', 'username', 'password']
    #         missing_fields = [f for f in required_fields if not data.get(f)]
    #         if missing_fields:
    #             raise ValueError("E603")
               

    #         # --- Chuẩn hóa ngày sinh (dob) về định dạng đúng
    #         try:
    #             dob_str = data.get('dob', '').strip()
    #             # Cho phép người dùng nhập cả 2 định dạng: 'yyyy-mm-dd' hoặc 'dd/mm/yyyy'
    #             if '/' in dob_str:
    #                 data['dob'] = datetime.strptime(dob_str, "%d/%m/%Y").date()
    #             else:
    #                 data['dob'] = datetime.strptime(dob_str, "%Y-%m-%d").date()
    #         except ValueError:
    #             raise ValueError("E603")
                

    #         # --- Kiểm tra mã sinh viên đã tồn tại
    #         existing = request.env['my_api_module.student'].search([('code', '=', data['code'])], limit=1)
    #         if existing:
    #             raise ValueError("E603")
               
           
    #         # --- Tạo sinh viên mới
    #         student = request.env['my_api_module.student'].sudo().create(data)

    #         return request.make_response(
    #             json.dumps({
    #                 "code": 200,
    #                 "status": "success",
    #                 "message": "Tạo mới sinh viên thành công",
    #                 "data": {
    #                     "id": student.id
    #                 }
    #             }),
    #             headers=[('Content-Type', 'application/json')]
    #         )
    #     except ValueError as ve:
    #         code  = str(ve)
    #         return request.make_response(
    #             json.dumps({
    #                 "code": code,
    #                 "status": "error",
    #                 "message": "Lỗi kiểm tra dữ liệu",
    #                 "data": {}
    #             }),
    #             headers=[('Content-Type', 'application/json')]
    #         )
    #     except Exception as e:
    #         return request.make_response(
    #             json.dumps({
    #                 "code": 600,
    #                 "status": "error",
    #                 "message": f"Lỗi không xác định: {str(e)}",
    #                 "data": {}
    #             }),
    #             headers=[('Content-Type', 'application/json')]
    #         )


##4: Lấy sinh viên theo ID
    @http.route('/tra_student/<int:id>', type='http', auth='public', methods=['GET'], csrf=False)
    def get_by_id(self, id, **kw):
        try:
            columnlist = kw.get('columnlist')

            # Bảng ánh xạ alias -> field thực tế
            alias_map = {
                'id': 'id',
                'co': 'code',
                'fu': 'fullname',
                'db': 'dob',
                'sx': 'sex',
                'hc': 'homecity',
                'ad': 'address',
                'ho': 'hobbies',
                'hc_': 'hair_color',
                'em': 'email',
                'fb': 'facebook',
                'cl': 'class_id',
                'us': 'username',
                'pw': 'password',
                'de': 'description',
                'at': 'attachment',
                'au': 'attachment_url'
            }

            # Nếu có columnlist thì dịch alias sang field thực tế
            if columnlist:
                aliases = columnlist.split(',')
                fields = []
                for a in aliases:
                    if a not in alias_map:
                        return request.make_response(json.dumps({
                            "code": "D607",
                            "status": "error",
                            "message": "Lỗi danh sách columnList",
                            "data": {}
                        }), headers=[('Content-Type', 'application/json')])
                    fields.append(alias_map[a])
            else:
                fields = list(alias_map.values())# lấy tất cả

            Student = request.env['my_api_module.student']
            student = Student.browse(id)

            if not student.exists():
                return request.make_response(json.dumps({
                    "code": "D604",
                    "status": "error",
                    "message": f"Không tìm thấy sinh viên với id = {id}",
                    "data": {}
                }), headers=[('Content-Type', 'application/json')])

            # Xây dựng dữ liệu trả về
            result = {}
            for field in fields :
                if not hasattr(student, field):
                    continue  # tránh lỗi nếu có field không tồn tại
                value = getattr(student, field)
                if isinstance(value, date):
                    result[field] = value.strftime('%Y-%m-%d')
                elif hasattr(value, 'id') and hasattr(value, 'name'):
                    result[field] = value.id if value else None  # Chỉ lấy ID nếu là Many2one
                elif field == 'attachment':
                    result[field] = ""
                elif field == 'hobbies':
                    result[field] = self.decode_bitmask_to_hobbies_string(value)
                else:
                    result[field] = value

            return request.make_response(json.dumps({
                "status": "success",
                "code": 200,
                # "message": f"Lấy thông tin sinh viên id = {id} thành công",
                "data": result
            }), headers=[('Content-Type', 'application/json')])

        except Exception as e:
            return request.make_response(json.dumps({
                "code": "D600",
                "status": "error",
                "message": f"Lỗi không xác định: {str(e)}",
                "data": {}
            }), headers=[('Content-Type', 'application/json')])


 

###5: CẬP NHẬT THÔNG TIN SINH VIÊN ( nếu thay đổi code mà không thay đổi id thì sẽ báo lỗi)
    @http.route('/tra_student/<int:id>', type='http', auth='public', methods=['PUT','POST'], csrf=False)
    def update(self, id, **kw):
        # _logger.info(f"[UPDATE FRONTEND] Request từ frontend: {kw}")
        # try:
        #     Student = request.env['my_api_module.student']
        #     student = Student.browse(id)

        #     if not student.exists():
        #         return self._json_error("F603", f"Không tìm thấy sinh viên với id = {id}")

        #     update_vals = {}
        #     allowed_fields = {
        #         'code', 'fullname', 'dob', 'sex', 'homecity', 'address',
        #         'hobbies', 'hair_color', 'email', 'facebook',
        #         'class_id', 'username', 'password', 'description', 'attachment'
        #     }
        #     _logger.info(f"[DEBUG] Các keys từ request.httprequest.files: {request.httprequest.files.keys()}")

        #     # Kiểm tra nếu có file
        #     files = request.httprequest.files.getlist('file')
        #     file = files[0] if files else None
        #     if file:
        #         filename = f"student_{id}_{file.filename}"
        #         filepath = os.path.join(UPLOAD_DIR, filename)
        #         with open(filepath, 'wb') as f:
        #             file.save(f)

        #         _logger.info(f"[FILE] Đã lưu file vào: {filepath}")

        #         # Đọc file CSV hoặc Excel để lấy dữ liệu cập nhật
        #         try:
        #             if filename.lower().endswith('.csv'):
        #                 df = pd.read_csv(filepath)
        #             elif filename.lower().endswith(('.xls', '.xlsx')):
        #                 df = pd.read_excel(filepath)
        #                 _logger.info(f"[DEBUG] DataFrame đọc được:\n{df}")
        #             else:
        #                 raise ValueError("Lỗi kiểm tra dữ liệu")

        #             if df.empty:
        #                 raise ValueError("Lỗi kiểm tra dữ liệu")

        #             record = df.iloc[0].to_dict()
        #             _logger.info(f"[PARSE FILE] Dữ liệu lấy từ file: {record}")

        #             for key, value in record.items():
        #                 if key not in allowed_fields:
        #                     continue
        #                 if pd.isna(value):
        #                     continue
        #                 if key == 'code':
        #                     duplicate = Student.search([('code', '=', value), ('id', '!=', id)], limit=1)
        #                     if duplicate:
        #                         raise ValueError(f"Mã sinh viên '{value}' đã tồn tại.")
        #                     update_vals['code'] = value

        #                 elif key == 'dob':
        #                     try:
        #                         parse_date = value.strftime('%Y-%m-%d') if isinstance(value, date) else str(value)
        #                         update_vals['dob'] = parse_date
        #                     except ValueError:
        #                         raise ValueError("Ngày sinh không hợp lệ. Định dạng hợp lệ: YYYY-MM-DD")

        #                 elif key == 'class_id':
        #                     try:
        #                         update_vals['class_id'] = int(value)
        #                     except Exception:
        #                         raise ValueError("class_id phải là số nguyên.")

        #                 elif key == 'hobbies':
        #                     try:
        #                         update_vals['hobbies'] = self.encode_hobbies_binary_string_to_bitmask(value)
        #                     except Exception:
        #                         raise ValueError("Hobbies không hợp lệ.")

        #                 else:
        #                     update_vals[key] = value

        #         except Exception as e:
        #             return request.make_response(json.dumps({
        #                 "code" : "F600",
        #                 "status": "error",
        #                 "message": f"Lỗi khi đọc file: {str(e)}",
        #                 "data": {}
        #             }), headers=[('Content-Type', 'application/json')])
        #         # update_vals['attachment'] = filename  # Ghi tên file lại nếu cần
        #     _logger.info(f"[UPDATE] Data cuối cùng: {update_vals}")
        #     # Xử lý các key gửi thủ công (text fields)
        #     for key, value in kw.items():
        #         if key not in allowed_fields:
        #             continue

        #         if key == 'code':
        #             duplicate = Student.search([('code', '=', value), ('id', '!=', id)], limit=1)
        #             if duplicate:
        #                 raise ValueError(f"Mã sinh viên '{value}' đã tồn tại.")
        #             update_vals['code'] = value

        #         elif key == 'dob':
        #             try:
        #                 parse_date = value.strftime('%Y-%m-%d') if isinstance(value, date) else str(value)
        #                 update_vals['dob'] = parse_date
        #             except ValueError:
        #                 raise ValueError("Ngày sinh không hợp lệ. Định dạng hợp lệ: YYYY-MM-DD")

        #         elif key == 'class_id':
        #             try:
        #                 update_vals['class_id'] = int(value)
        #             except Exception:
        #                 raise ValueError("class_id phải là số nguyên.")

        #         elif key == 'hobbies':
        #             try:
        #                 update_vals['hobbies'] = self.encode_hobbies_binary_string_to_bitmask(value)
        #             except Exception:
        #                 raise ValueError("Hobbies không hợp lệ.")

        #         else:
        #             update_vals[key] = value

            
        #     _logger.info(f"[DEBUG] student type = {type(student)}")  # Xem chính xác là gì

        #     student.write(update_vals)
        #     return request.make_response(json.dumps({
        #         "code": "200",
        #         "status": "success",
        #         "message": f"Cập nhật sinh viên id = {id} thành công",
        #         "data": {"id": student.id}
        #     }), headers=[('Content-Type', 'application/json')])

        # except ValueError as ve:
        #     # Lỗi kiểm tra dữ liệu
        #     return request.make_response(json.dumps({
        #         "code": "F603",
        #         "status": "error",
        #         "message": str(ve),
        #         "data": {}
        #     }), headers=[('Content-Type', 'application/json')])

        # except Exception as e:
        #     # Lỗi không xác định
        #     return request.make_response(json.dumps({
        #         "code": "F600",
        #         "status": "error",
        #         "message": f"Lỗi không xác định: {str(e)}",
        #         "data": {}
        #     }), headers=[('Content-Type', 'application/json')])
        _logger.info(f"[UPDATE FRONTEND] Request từ frontend: {kw}")
        
        try:
            Student = request.env['my_api_module.student']
            student = Student.browse(id)

            if not student.exists():
                raise ValueError(f"Không tìm thấy sinh viên với id = {id}")

            update_vals = {}
            allowed_fields = {
                'code', 'fullname', 'dob', 'sex', 'homecity', 'address',
                'hobbies', 'hair_color', 'email', 'facebook',
                'class_id', 'username', 'password', 'fattachment', 'attachment_url'
            }
            
            # Khởi tạo update_vals với giá trị rỗng cho các trường Char/Text và False cho Many2one
            
            # Kết hợp dữ liệu từ file và kw.items()
            data = kw # Bắt đầu với dữ liệu từ kw
            _logger.info(f"[DEBUG] Các keys từ request.httprequest.files: {request.httprequest.files.keys()}")

            # Xử lý file cập nhật thông tin nếu có
            files = request.httprequest.files.getlist('file')
            file = files[0] if files else None
            timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
            if file:
                filename = f"student_{timestamp}_{student.id}_{file.filename}"
                filepath = os.path.join(DOCS_DIR, filename)
                # Xử lý xóa file cũ của id nếu có
                old_files_pattern = os.path.join(DOCS_DIR, f"student_*_{student.id}_*")
                old_files = glob.glob(old_files_pattern)
                for old_file in old_files:
                    os.remove(old_file)
                    _logger.info(f"[FILE] Đã xóa file cũ: {old_file}")
                with open(filepath, 'wb') as f:
                    file.save(f)
                _logger.info(f"[FILE] Đã lưu file vào: {filepath}")

                # Đọc file CSV hoặc Excel
                try:
                    if filename.lower().endswith('.csv'):
                        df = pd.read_csv(filepath)
                        if 'dob' in df.columns:
                            df['dob'] = pd.to_datetime(df['dob']).dt.strftime('%Y-%m-%d')
                    elif filename.lower().endswith(('.xls', '.xlsx')):
                        df = pd.read_excel(filepath)
                        if 'dob' in df.columns:
                            df['dob'] = pd.to_datetime(df['dob']).dt.strftime('%Y-%m-%d')
                        _logger.info(f"[DEBUG] DataFrame đọc được:\n{df}")
                    else:
                        raise ValueError("Lỗi kiểm tra dữ liệu")

                    if df.empty:
                        raise ValueError("Lỗi kiểm tra dữ liệu")

                    record = df.iloc[0].to_dict()
                    _logger.info(f"[PARSE FILE] Dữ liệu lấy từ file: {record}")
                    # Gộp dữ liệu từ file vào data, ưu tiên dữ liệu từ kw nếu trùng key
                    for key, value in record.items():
                        if key not in allowed_fields or pd.isna(value):
                            continue
                        data[key] = value  # Ghi đè hoặc thêm vào data

                except Exception as e:
                    return request.make_response(json.dumps({
                        "code": "F600",
                        "status": "error",
                        "message": f"Lỗi khi đọc file: {str(e)}",
                        "data": {}
                    }), headers=[('Content-Type', 'application/json')])
            validator = StudentValidator(data, request.env['my_api_module.student'])
            errors = validator.validate_update_data(id)
            if errors:
                _logger.info(f"[UPDATE] Lỗi kiểm tra dữ liệu: {errors}")
                raise ValueError(" ".join(errors))
            _logger.info(f"định dạng của trường dob:{type(data.get('dob'))}")
           
            # Xử lý dữ liệu hợp nhất
            for key, value in data.items():
                if key not in allowed_fields:
                    continue

               
                elif key == 'hobbies':
                    try:
                        update_vals['hobbies'] = self.encode_hobbies_binary_string_to_bitmask(value)
                    except Exception as e:
                        raise ValueError("Lỗi encode hobbies.")    
                elif key == 'fattachment':
     # Xử lý image
                    
                    file_image = request.httprequest.files.get("fattachment")
                    _logger.info("[STORE] Danh sách file upload: %s", list(request.httprequest.files.keys()))
                    if file_image and hasattr(file_image, 'filename') and file_image.filename:
                        _logger.info("[STORE] Nhận được file: %s", file_image.filename)
                        # Tạo tên file an toàn và duy nhất
                        filename = f"student_{timestamp}_{student.id}_{secure_filename(file_image.filename)}" #trả về định dạng an toàn 'My_cool_movie.mov'
                        filepath = os.path.join(IMAGES_DIR, filename)
                        try:
                            # Đảm bảo thư mục IMAGES_DIR tồn tại và có quyền ghi
                            if not os.access(IMAGES_DIR, os.W_OK):
                                _logger.error(f"[FATTACHMENT] Thư mục {IMAGES_DIR} không có quyền ghi")
                                raise Exception(f"Thư mục {IMAGES_DIR} không có quyền ghi")
                            # os.makedirs(IMAGES_DIR, exist_ok=True)
                            # Xóa file ảnh cũ của id nếu có
                            old_image_pattern = os.path.join(IMAGES_DIR, f"student_*_{student.id}_*")
                            old_images = glob.glob(old_image_pattern)
                            for old_image_file in old_images:
                                os.remove(old_image_file)
                                _logger.info(f"[FATTACHMENT] Đã xóa ảnh cũ: {old_image_file}")
                                
                            # Ghi file vào thư mục
                            with open(filepath, 'wb') as f:
                                file_image.save(f)
                            # Tạo URL đầy đủ
                            attachment_url = f"{BASE_URL}/images/{filename}"
                            attachment = f"images/{filename}"
                            update_vals['attachment_url'] = attachment_url
                            update_vals['attachment'] = attachment
                            _logger.info(f"[FATTACHMENT] File uploaded to: {attachment_url}")
                            _logger.info(f"[FATTACHMENT] File saved to: {filepath}")
                        except Exception as e:
                            _logger.error(f"[FATTACHMENT] Lỗi khi ghi file: {str(e)}")
                            raise ValueError(f"Lỗi khi ghi file: {str(e)}")
                elif 'fattachment' not in kw and 'attachment' not in kw:
                    old_image_pattern = os.path.join(IMAGES_DIR, f"student_*_{student.id}_*")
                    old_images = glob.glob(old_image_pattern)
                    for old_image_file in old_images:
                        os.remove(old_image_file)
                        _logger.info(f"[FATTACHMENT] Đã xóa ảnh cũ: {old_image_file}")
                        
                    update_vals['attachment_url'] = ""
                    update_vals['attachment'] = ""
                else:
                    update_vals[key] = "" if value is None or False else value
            
            _logger.info("[UPDATE] Dữ liệu sau khi phân tích: %s", update_vals)
            #     elif key == 'fattachment':
            #                 # Xử lý image
            #         if value is None or False:
            #             # Xóa ảnh cũ của student
            #             old_image_pattern = os.path.join(IMAGES_DIR, f"student_{student.id}_*")
            #             old_images = glob.glob(old_image_pattern)
            #             for old_image_file in old_images:
            #                 try:
            #                     os.remove(old_image_file)
            #                     _logger.info(f"[FATTACHMENT] Đã xóa ảnh cũ: {old_image_file}")
            #                 except Exception as e:
            #                     _logger.error(f"[FATTACHMENT] Lỗi khi xóa ảnh cũ {old_image_file}: {str(e)}")
            #             # Xóa luôn attachment_url trong DB
            #             update_vals['attachment_url'] = ""
            #             update_vals['fattachment'] = ""  # Đặt lại fattachment về rỗng
            #             _logger.info("[FATTACHMENT] Đã xóa ảnh cũ và đặt attachment_url về rỗng")
            #         else:  # Chỉ xử lý nếu value không rỗng
            #             timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
            #             file_image = request.httprequest.files.get("fattachment")
            #             _logger.info("[STORE] Danh sách file upload: %s", list(request.httprequest.files.keys()))
            #             if file_image and hasattr(file_image, 'filename') and file_image.filename:
            #                 _logger.info("[STORE] Nhận được file: %s", file_image.filename)
            #                 # Tạo tên file an toàn và duy nhất
            #                 filename = f"student_{timestamp}_{student.id}_{secure_filename(file_image.filename)}"
            #                 filepath = os.path.join(IMAGES_DIR, filename)
            #                 try:
            #                     # Đảm bảo thư mục IMAGES_DIR tồn tại và có quyền ghi
            #                     if not os.access(IMAGES_DIR, os.W_OK):
            #                         _logger.error(f"[FATTACHMENT] Thư mục {IMAGES_DIR} không có quyền ghi")
            #                         raise Exception(f"Thư mục {IMAGES_DIR} không có quyền ghi")
            #                     os.makedirs(IMAGES_DIR, exist_ok=True)
            #                     # Xử lý xóa file ảnh cũ của id nếu có
            #                     old_image_pattern = os.path.join(IMAGES_DIR, f"student_*_{student.id}_*")
            #                     old_images = glob.glob(old_image_pattern)
            #                     for old_image_file in old_images:
            #                         try:
            #                             os.remove(old_image_file)
            #                             _logger.info(f"[FATTACHMENT] Đã xóa ảnh cũ: {old_image_file}")
            #                         except Exception as e:
            #                             _logger.error(f"[FATTACHMENT] Lỗi khi xóa ảnh cũ {old_image_file}: {str(e)}")
            #                     # Ghi file vào thư mục
            #                     with open(filepath, 'wb') as f:
            #                         file_image.save(f)
            #                     # Tạo URL đầy đủ
            #                     attachment_url = f"{BASE_URL}/images/{filename}"
            #                     attachment = f"images/{filename}"
            #                     update_vals['attachment_url'] = attachment_url
            #                     update_vals['attachment'] = attachment
            #                     # update_vals['fattachment'] = filepath
            #                     _logger.info(f"[FATTACHMENT] File uploaded to: {attachment_url}")
            #                     _logger.info(f"[FATTACHMENT] File saved to: {filepath}")
            #                 except Exception as e:
            #                     _logger.error(f"[FATTACHMENT] Lỗi khi ghi file: {str(e)}")
            #                     raise ValueError(f"Lỗi khi ghi file: {str(e)}")
            #          # Không có file, để rỗng
            #     else:
            #         update_vals[key] = "" if value is None or False else value
                
                
            # if 'fattachment' not in data and 'attachment' not in data:
            #     update_vals['attachment_url'] = ""
            #     # Xóa ảnh cũ của student
            #     old_image_pattern = os.path.join(IMAGES_DIR, f"student_*_{id}_*")
            #     old_images = glob.glob(old_image_pattern)
            #     for old_image_file in old_images:
            #         try:
            #             os.remove(old_image_file)
            #             _logger.info(f"[UPDATE] Đã xóa ảnh cũ của attachment: {old_image_file}")
            #         except Exception as e:
            #             _logger.error(f"[UPDATE] Lỗi khi xóa ảnh cũ {old_image_file}: {str(e)}")
            # _logger.info("[UPDATE] Dữ liệu sau khi phân tích: %s", update_vals)
                
            student.write(update_vals)
            return request.make_response(json.dumps({
                "code": "200",
                "status": "success",
                "message": f"Cập nhật sinh viên id = {id} thành công",
                "data": {"id": student.id}
            }), headers=[('Content-Type', 'application/json')])

        except ValueError as ve:
            return request.make_response(json.dumps({
                "code": "F603",
                "status": "error",
                "message": f"Lỗi kiểm tra dữ liệu: {str(ve)}" ,
                "data": {}
            }), headers=[('Content-Type', 'application/json')])

        except Exception as e:
            return request.make_response(json.dumps({
                "code": "F600",
                "status": "error",
                "message": f"Lỗi không xác định: {str(e)}",
                "data": {}
            }), headers=[('Content-Type', 'application/json')])
##6: XÓA SINH VIÊN THEO ID

    @http.route('/tra_student/<int:id>', type='http', auth='public', methods=['DELETE'], csrf=False)
    def destroy(self, id, **kw):
        try:
            Student = request.env['my_api_module.student']
            student = Student.browse(id)

            # Kiểm tra tồn tại
            if not student.exists():
                return request.make_response(json.dumps({
                    "code": "G604",
                    "status": "error",
                    "message": f"Không tìm thấy sinh viên với id = {id}",
                    "data": {}
                }), headers=[('Content-Type', 'application/json')])
            for directory in [IMAGES_DIR, DOCS_DIR]:
                file_pattern = os.path.join(directory, f"student_*_{student.id}_*")
                files = glob.glob(file_pattern)
                for file_path in files:
                    os.remove(file_path)
                    _logger.info(f"[DELETE] Đã xóa file: {file_path}")
                       
            try:
                student.unlink()
            except Exception as e:
                # Xử lý lỗi liên kết khóa ngoại
                if 'foreign key' in str(e).lower():
                    return request.make_response(json.dumps({
                        "code": "G605",
                        "status": "error",
                        "message": f"Bản ghi id = {id} đang được tham chiếu và không thể xóa",
                        "data": {}
                    }), headers=[('Content-Type', 'application/json')])
                else:
                    raise e  # Không phải lỗi khóa ngoại, ném lại để xử lý dưới

            return request.make_response(json.dumps({
                "code": "200",
                "status": "success",
                "message": f"Đã xóa sinh viên có id = {id}",
                "data": {
                    "id": id
                }
            }), headers=[('Content-Type', 'application/json')])

        except Exception as e:
            return request.make_response(json.dumps({
                "code": "G600",
                "status": "error",
                "message": f"Lỗi không xác định: {str(e)}",
                "data": {}
            }), headers=[('Content-Type', 'application/json')])
            
##7: Tạo bản sao của sinh viên theo ID
    @http.route('/tra_student/copy/<int:id>', type='http', auth='public', methods=['POST'], csrf=False)
    def copy(self, id):
        try:
            # Tìm bản ghi gốc
            student = request.env['my_api_module.student'].browse(id)

            # Kiểm tra bản ghi có tồn tại không
            if not student.exists():
                return request.make_response(
                    json.dumps({
                        "code": "H604",
                        "status": "error",
                        "message": "Sinh viên không tồn tại",
                        "data": {}
                    }),
                    headers=[('Content-Type', 'application/json')]
                )

            # Kiểm tra dữ liệu cơ bản
            if not student.code or not student.fullname:
                return request.make_response(
                    json.dumps({
                        "code": "H603",
                        "status": "error",
                        "message": "Thiếu mã sinh viên hoặc họ tên",
                        "data": {}
                    }),
                    headers=[('Content-Type', 'application/json')]
                )

            # Gọi phương thức copy (duplicate), tạm thời giữ nguyên code gốc
            copied_student = student.copy()

            # Tạo code mới: gốc + _id (VD: SV407_36)
            new_code = f"{student.code}_{copied_student.id}"
            copied_student.write({'code': new_code})

            # Trả về dữ liệu bản sao
            result_data = {
                'id': copied_student.id,
                'code': copied_student.code,
                'fullname': copied_student.fullname,
                'dob': copied_student.dob.strftime('%Y-%m-%d') if copied_student.dob else None,
                'sex': copied_student.sex,
                'homecity': copied_student.homecity,
                'address': copied_student.address,
                'hobbies': copied_student.hobbies,
                'hair_color': copied_student.hair_color,
                'email': copied_student.email,
                'facebook': copied_student.facebook,
                'class_id': copied_student.class_id.id if copied_student.class_id else None,
                'username': copied_student.username,
                'password': copied_student.password,
                'description': copied_student.description,
                'attachment': copied_student.attachment,
                'attachment_url': copied_student.attachment_url
            }

            return request.make_response(
                json.dumps({
                    "code": 200,
                    "status": "success",
                    "message": "Sao chép sinh viên thành công",
                    "data": result_data
                }),
                headers=[('Content-Type', 'application/json')]
            )

        except Exception as e:
            return request.make_response(
                json.dumps({
                    "code": "H600",
                    "status": "error",
                    "message": f"Lỗi không xác định: {str(e)}",
                    "data": {}
                }),
                headers=[('Content-Type', 'application/json')]
            )
##8: Tạo bản sao nhiều sinh viên   
    @http.route('/tra_student/copy', type='http', auth='public', methods=['POST'], csrf=False)
    def mass_copy(self, **kw):
        try:
            idlist_str = kw.get('idlist', '')
            if not idlist_str:
                return request.make_response(
                    json.dumps({
                        'code': 'H603',
                        'status': 'error',
                        'message': 'Thiếu danh sách ID',
                        'data': []
                    }),
                    headers=[('Content-Type', 'application/json')]
                )

            idlist = [int(i.strip()) for i in idlist_str.split(',') if i.strip().isdigit()]
            if not idlist:
                return request.make_response(
                    json.dumps({
                        'code': 'H603',
                        'status': 'error',
                        'message': 'Danh sách ID không hợp lệ',
                        'data': []
                    }),
                    headers=[('Content-Type', 'application/json')]
                )

###Trả về records matching search criteria
            students = request.env['my_api_module.student'].search([('id', 'in', idlist)]) 
            existing_ids = set(students.ids)
            missing_ids = [i for i in idlist if i not in existing_ids]

            if missing_ids:
                return request.make_response(
                    json.dumps({
                        'code': 'H604',
                        'status': 'error',
                        'message': f"Các ID không tồn tại: {missing_ids}",
                        'data': []
                    }),
                    headers=[('Content-Type', 'application/json')]
                )

            copied_students = []
            for student in students:
                copied = student.copy()
                new_code = f"{student.code}_{copied.id}"
                copied.write({'code': new_code})
                copied_students.append({
                    'id': copied.id,
                    'code': new_code,
                    'fullname': copied.fullname,
                    'dob': copied.dob.strftime('%Y-%m-%d') if copied.dob else None,
                    'sex': copied.sex,
                    'homecity': copied.homecity,
                    'address': copied.address,
                    'hobbies': copied.hobbies,
                    'hair_color': copied.hair_color,
                    'email': copied.email,
                    'facebook': copied.facebook,
                    'class_id': copied.class_id.id if copied.class_id else None,
                    'username': copied.username,
                    'password': copied.password,
                    'description': copied.description,
                    'attachment': copied.attachment,
                    'attachment_url': copied.attachment_url
                })

            return request.make_response(
                json.dumps({
                    'code': '200',
                    'status': 'success',
                    'message': 'Sao chép thành công',
                    'data': copied_students
                }),
                headers=[('Content-Type', 'application/json')]
            )

        except Exception as e:
            # _logger.error("Lỗi trong quá trình sao chép sinh viên: %s", str(e))
            return request.make_response(
                json.dumps({
                    'code': 'H600',
                    'status': 'error',
                    'message': f'Lỗi không xác định: {str(e)}',
                    'data': []
                }),
                headers=[('Content-Type', 'application/json')]
            )

##9: XÓA NHIỀU LỚP
    @http.route('/tra_student/delete', type='http', auth='public', methods=['DELETE'], csrf=False)
    def mass_delete(self, **kw):
        try:
            idlist_raw = kw.get('idlist')
            if not idlist_raw:
                raise ValueError('I604')  # Thiếu tham số idlist

            idlist = [int(i.strip()) for i in idlist_raw.split(',') if i.strip().isdigit()]
            if not idlist:
                raise ValueError('I604')  # Danh sách rỗng hoặc sai định dạng

            Student = request.env['my_api_module.student']
            records = Student.browse(idlist).exists()  # <-- Chỉ giữ bản ghi tồn tại thật sự

            existing_ids = records.mapped('id') ## Trả về list các id của bản ghi tồn tại
            not_found_ids = [i for i in idlist if i not in existing_ids]
            if not_found_ids:
                raise ValueError('I604')
            
            # Xử lý xóa file đính kèm trước
            for student in records:
                for directory in [IMAGES_DIR, DOCS_DIR]:
                    file_pattern = os.path.join(directory, f"student_*_{student.id}_*")
                    files = glob.glob(file_pattern)
                    for file_path in files:
                        os.remove(file_path)
                              
            #Xóa bản ghi
            try:
                records.unlink()
            except Exception:
                raise ValueError('I604')  # Có thể do ràng buộc khóa ngoại

            return request.make_response(
                json.dumps({
                    "code": 200,
                    "status": "success",
                    "message": "Xóa các bản ghi thành công",
                    "data": {"ids": existing_ids}
                }),
                headers=[('Content-Type', 'application/json')]
            )

        except ValueError as ve:
            return request.make_response(
                json.dumps({
                    "code": str(ve),
                    "status": "error",
                    "message": "Lỗi các bản ghi không xóa được.",
                    "data": {}
                }),
                headers=[('Content-Type', 'application/json')]
            )

        except Exception as e:
            return request.make_response(
                json.dumps({
                    "code": "I600",
                    "status": "error",
                    "message": f"Lỗi không xác định: {str(e)}",
                    "data": {}
                }),
                headers=[('Content-Type', 'application/json')]
            )

##10: import dữ liệu từ file CSV
    @http.route('/tra_student/import', type='http', auth='public', methods=['POST'], csrf=False)
    def import_data(self, **kw):
        try:
            # ==== Kiểm tra file đính kèm ====
            file = kw.get('attachment')
            if not file:
                raise ValueError("J604")

            filename = file.filename.lower()
            if not (filename.endswith('.csv') or filename.endswith('.xlsx') or filename.endswith('.xls')):
                raise ValueError("J601")

            # ==== Đọc dữ liệu từ file ====
            records = []
            if filename.endswith('.csv'):
                content = file.stream.read().decode('utf-8')
                reader = csv.DictReader(io.StringIO(content))
                for row in reader:
                    records.append(row)

            elif filename.endswith(('.xlsx', '.xls')):
                if not pd:
                    raise ImportError("Chưa cài pandas hoặc openpyxl.")
                df = pd.read_excel(file.stream)
                records = df.to_dict(orient='records')

            if not records:
                raise ValueError("J605")

            # ==== Thêm sinh viên vào hệ thống ====
            created = []
            for row in records:
                code = row.get('code')
                fullname = row.get('fullname')
                dob = row.get('dob')
                if 'sex' in row:
                    row['sex'] = str(row['sex']) if row['sex'] is not None else '1'
                sex = row.get('sex')
                homecity = row.get('homecity')
                address = row.get('address')
                hobbies = row.get('hobbies')
                hair_color = row.get('hair_color')
                email = row.get('email')
                facebook = row.get('facebook')
                class_id = row.get('class_id')
                username = row.get('username')
                password = row.get('password')
                description = row.get('description')
                attachment = row.get('attachment')

                if not code or not fullname:
                    continue  # Bỏ qua bản ghi thiếu thông tin tối thiểu

                # Kiểm tra xem code đã tồn tại chưa
                existing = request.env['my_api_module.student'].search([('code', '=', code)], limit=1)
                temp_code = code if not existing else ''

                # Tạo bản ghi mới
                new_student = request.env['my_api_module.student'].create({
                    'code': temp_code,
                    'fullname': fullname,
                    'dob': dob,
                    'sex': sex,
                    'homecity': homecity,
                    'address': address,
                    'hobbies': hobbies,
                    'hair_color': hair_color,
                    'email': email,
                    'facebook': facebook,
                    'class_id': int(class_id) if class_id else None,
                    'username': username,
                    'password': password,
                    'description': description,
                    'attachment': attachment
                })

                # Nếu code bị trùng, cập nhật lại theo dạng SV407_<id>
                if not temp_code:
                    new_code = f"{code}_{new_student.id}"
                    new_student.write({'code': new_code})
                else:
                    new_code = code

                created.append({
                    'id': new_student.id,
                    'code': new_code,
                    'fullname': new_student.fullname,
                    'dob': new_student.dob.strftime('%Y-%m-%d') if new_student.dob else None,
                    'sex': new_student.sex,
                    'homecity': new_student.homecity,
                    'address': new_student.address,
                    'hobbies': new_student.hobbies,
                    'hair_color': new_student.hair_color,
                    'email': new_student.email,
                    'facebook': new_student.facebook,
                    'class_id': new_student.class_id.id if new_student.class_id else None,
                    'username': new_student.username,
                    'password': new_student.password,
                    'description': new_student.description,
                    'attachment': new_student.attachment
                })

            if not created:
                raise ValueError("J605")

            # ==== Trả kết quả thành công ====
            return request.make_response(
                json.dumps({
                    "code": "200",
                    "status": "success",
                    "message": "Import thành công",
                    "data": created
                }, ensure_ascii=False),
                headers=[('Content-Type', 'application/json')]
            )

        except ValueError as ve:
            code = str(ve)
            messages = {
                'J601': "Lỗi định dạng tệp. Chỉ hỗ trợ CSV hoặc Excel.",
                'J604': "Lỗi tệp dữ liệu chưa tải lên.",
                'J605': "Lỗi không có dữ liệu được thêm mới từ file."
            }
            return request.make_response(
                json.dumps({
                    "code": code,
                    "status": "error",
                    "message": messages.get(code, f"Lỗi: {code}"),
                    "data": []
                }, ensure_ascii=False),
                headers=[('Content-Type', 'application/json')]
            )

        except Exception as e:
            return request.make_response(
                json.dumps({
                    "code": "J600",
                    "status": "error",
                    "message": f"Lỗi không xác định: {str(e)}",
                    "data": []
                }, ensure_ascii=False),
                headers=[('Content-Type', 'application/json')]
            )

##11: Xuất BẢN GHI THEO ID
    @http.route('/tra_student/export/<int:record_id>', type='http', auth='public', methods=['GET'], csrf=False)
    def export_by_id(self, record_id, **kw):
        try:
            export_type = kw.get('type', 'csv').lower()
            columnlist = kw.get('columnlist')

            if not columnlist:
                raise ValueError("K607")

            aliases = [a.strip() for a in columnlist.split(',')]

            alias_to_field = {
                'id': 'id',
                'co': 'code',
                'fu': 'fullname',
                'db': 'dob',
                'sx': 'sex',
                'hc': 'homecity',
                'ad': 'address',
                'ho': 'hobbies',
                'hc_': 'hair_color',
                'em': 'email',
                'fb': 'facebook',
                'cl': 'class_id',
                'us': 'username',
                'pw': 'password',
                'de': 'description',
                'at': 'attachment',
            }

            fields = []
            for alias in aliases:
                if alias not in alias_to_field:
                    raise ValueError(f"Bí danh '{alias}' không hợp lệ.")
                fields.append(alias_to_field[alias])

            if export_type not in ['csv', 'xlsx', 'docx']:
                raise ValueError("K601")

            student_model = request.env['my_api_module.student']
            record = student_model.browse(record_id)

            if not record.exists():
                raise ValueError("K604")

            row = []
            for field in fields:
                value = getattr(record, field, '')
                if field == 'dob' and value:
                    value = value.strftime('%Y-%m-%d')
                elif field == 'class_id':
                    value = value.id if value else None
                elif isinstance(value, (list, tuple)):
                    value = ', '.join(str(v) for v in value) # chuyển đổi danh sách sang chuỗi
                elif value is False or value is None:
                    value = ''
                row.append(value)

            # Xử lý xuất tệp 
            data_dict = [dict(zip(fields, row))]

            if export_type == 'csv':
                buffer = io.StringIO()
                pd.DataFrame(data_dict).to_csv(buffer, index=False)
                response_data = buffer.getvalue().encode('utf-8')
                content_type = 'text/csv'

            elif export_type == 'xlsx':
                output = io.BytesIO()
                pd.DataFrame(data_dict).to_excel(output, index=False) #	Biến chuỗi (str) thành bytes để gửi qua HTTP
                response_data = output.getvalue()
                content_type = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'

            elif export_type == 'docx':
                doc = Document()
                doc.add_heading('Thông tin sinh viên', 0)
                table = doc.add_table(rows=1, cols=len(fields))
                hdr_cells = table.rows[0].cells ## Lẩy ra các ô ở hàng đầu tiên
                for idx, field in enumerate(fields):
                    hdr_cells[idx].text = field 
                row_cells = table.add_row().cells ## Thêm một hàng
                for idx, cell in enumerate(row):
                    row_cells[idx].text = str(cell)
                output = io.BytesIO()
                doc.save(output)
                response_data = output.getvalue()
                content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

            return request.make_response(
                response_data,
                headers=[
                    ('Content-Type', content_type),
                    ('Content-Disposition', f'attachment; filename=student_{record_id}.{export_type}')
                ]
            )

        except ValueError as ve:
            code = str(ve)
            messages = {
                'K601': "Lỗi định dạng tệp. Chỉ hỗ trợ CSV, XLSX hoặc DOCX.",
                'K604': "Lỗi mã bản ghi không tồn tại",
                'K607': "Lỗi danh sách columnList."
            }
            return request.make_response(
                json.dumps({
                    "code": code,
                    "status": "error",
                    "message": messages.get(code, f"Lỗi: {code}")
                }),
                headers=[('Content-Type', 'application/json')]
            )

        except Exception as e:
            return request.make_response(
                json.dumps({
                    "code": "K600",
                    "status": "error",
                    "message": f"Lỗi không xác định: {str(e)}"
                }),
                headers=[('Content-Type', 'application/json')]
            )


##12: xuất nhiều bản ghi
    @http.route('/tra_student/export', type='http', auth='public', methods=['POST'], csrf=False)
    def mass_export(self, **kw):
        try:
            idlist_raw = kw.get('idlist')
            export_type = kw.get('type', 'csv').lower()
            columnlist = kw.get('columnlist')

            if not idlist_raw or not columnlist:
                raise ValueError("L607")

            try:
                idlist = [int(x) for x in idlist_raw.split(',') if x.strip().isdigit()]
            except Exception:
                raise ValueError("L607")

            if not idlist:
                raise ValueError("L607")

            aliases = [a.strip() for a in columnlist.split(',')]

            alias_to_field = {
                'id': 'id',
                'co': 'code',
                'fu': 'fullname',
                'db': 'dob',
                'sx': 'sex',
                'hc': 'homecity',
                'ad': 'address',
                'ho': 'hobbies',
                'hc_': 'hair_color',
                'em': 'email',
                'fb': 'facebook',
                'cl': 'class_id',
                'us': 'username',
                'pw': 'password',
                'de': 'description',
                'at': 'attachment',
            }

            fields = []
            for alias in aliases:
                if alias not in alias_to_field:
                    raise ValueError("L607")
                fields.append(alias_to_field[alias])

            if export_type not in ['csv', 'xlsx', 'docx']:
                raise ValueError("L601")

            student_model = request.env['my_api_module.student']
            #### filtered : trả về các bản ghi thỏa mãn func, exists: trả về các bản ghi tồn tại
            records = student_model.browse(idlist).filtered(lambda r: r.exists())

            if not records or len(records) == 0 or len(records) != len(idlist):
                raise ValueError("L602")

            data_list = []
            for record in records:
                row = []
                for field in fields:
                    value = getattr(record, field, '')
                    if field == 'dob' and value:
                        value = value.strftime('%Y-%m-%d')
                    elif field == 'class_id':
                        value = value.id if value else None
                    # elif isinstance(value, (list, tuple)):
                    #     value = ', '.join(str(v) for v in value)
                    elif value is False or value is None:
                        value = ''
                    row.append(value)
                data_list.append(dict(zip(fields, row)))  # Trả về một list các dict

            # Export to selected file type
            if export_type == 'csv':
                buffer = io.StringIO()
                pd.DataFrame(data_list).to_csv(buffer, index=False)
                response_data = buffer.getvalue()
                content_type = 'text/csv'

            elif export_type == 'xlsx':
                output = io.BytesIO()
                pd.DataFrame(data_list).to_excel(output, index=False)
                response_data = output.getvalue()
                content_type = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'

            elif export_type == 'docx':
                doc = Document()
                doc.add_heading('Thông tin nhiều sinh viên', 0)
                table = doc.add_table(rows=1, cols=len(fields))
                hdr_cells = table.rows[0].cells # Lấy ra các ô ở hàng đầu tiên
                for i, field in enumerate(fields):
                    hdr_cells[i].text = field # Lấy tên các trường làm tiêu đề cột
                for row_data in data_list:
                    row_cells = table.add_row().cells  
                    for i, val in enumerate(row_data.values()):
                        row_cells[i].text = str(val) # Ghi dữ liệu vào ô mới 
                output = io.BytesIO() # Tạo vùng nhớ đệm để ghi file 
                doc.save(output)
                response_data = output.getvalue()
                content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

            return request.make_response(
                response_data,
                headers=[
                    ('Content-Type', content_type),
                    ('Content-Disposition', f'attachment; filename=students.{export_type}')
                ]
            )

        except ValueError as ve:
            code = str(ve)
            messages = {
                'L601': "Lỗi định dạng tệp. Chỉ hỗ trợ CSV, XLSX hoặc DOCX.",
                'L602': "Danh sách bản ghi không tồn tại.",
                'L607': "Danh sách columnlist không hợp lệ."
            }
            return request.make_response(
                json.dumps({
                    "code": code,
                    "status": "error",
                    "message": messages.get(code, f"Lỗi: {code}")
                }),
                headers=[('Content-Type', 'application/json')]
            )
        except Exception as e:
            return request.make_response(
                json.dumps({
                    "code": "L600",
                    "status": "error",
                    "message": f"Lỗi không xác định: {str(e)}"
                }),
                headers=[('Content-Type', 'application/json')]
            )
