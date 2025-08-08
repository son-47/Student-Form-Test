from odoo import http
from odoo.http import request,Response, content_disposition
import json,io
import os,glob 
import re
import base64
import shutil, uuid
import mimetypes 
from datetime import datetime, date
import logging
from docx import Document
import pandas as pd
from werkzeug.utils import secure_filename
_logger = logging.getLogger(__name__)
from ..validator.class_validator import ClassValidator

MODULE_DIR = os.path.dirname(os.path.abspath(__file__))  # /addons/my_api_module/controllers/
BASE_DIR = os.path.abspath(os.path.join(MODULE_DIR, '..'))  # dẫn lên thư mục cha /addons/my_api_module/ 
DOCS_DIR = os.path.join(BASE_DIR, 'static', 'uploads', 'docs')  # Đường dẫn tới static/uploads/docs


## type : lấy dữ liệu đầu vào từ đâu
## auth : thiết lập quyền truy cập
########: # public: không cần đăng nhập
    #         # user: cần đăng nhập
    #         # 'bearer': (API bảo mật nâng cao)
    #         # 'none' : Không liên quan gì đến user hay database. 
    #           Dùng trong các tình huống hệ thống chưa khởi tạo database  
    
## csrf : Cross-Site Request Forgery, bảo vệ chống lại các cuộc tấn công giả mạo yêu cầu từ trang web khác
# khi tắt thì Cho phép gửi request từ client khác (Postman, mobile app)

##1 : LẤY TOÀN BỘ LỚP
class ClassController(http.Controller):
    @http.route('/tra_class', type='http', auth='public', methods=['GET'], csrf=False)
    def get_all(self, **kwargs): 
        try:
            columnlist = kwargs.get('columnlist')
            fields_list = columnlist.split(',') if columnlist else ['id', 'code', 'name', 'description', 'write_date']

            classes = request.env['my_api_module.classes'].search([])

            data = []
            for c in classes:
                record = {}
                for field in fields_list:
                    try:
                        value = getattr(c, field)
                        # Nếu là datetime → chuyển sang chuỗi định dạng ISO hoặc định dạng đơn giản
                        if isinstance(value, datetime):
                            value = value.strftime('%Y-%m-%d')
                        record[field] = value if value is not None else ''
                    except Exception:
                        record[field] = ''
                data.append(record)

            return request.make_response(
                json.dumps({
                    "code": 200,
                    "status": "success",
                    "message": "Lấy danh sách lớp thành công",
                    "data": data
                }),
                headers=[('Content-Type', 'application/json')]
            )

        except Exception as e:
            return request.make_response(
                json.dumps({
                    "code": "B600",
                    "status": "error",
                    "message": f"Lỗi không xác định: {str(e)}",
                    "data": []
                }),
                headers=[('Content-Type', 'application/json')]
            )



###2 : LẤY THEO PAGE
    @http.route('/tra_class/page/<int:init>', type='http', auth='public', methods=['GET'], csrf=False)
    def get_by_page(self, init, **kwargs):
        try:
            # Kiểm tra số trang
            if init < 1:
                raise ValueError('C601')  # Số trang không hợp lệ

            # Kích thước trang
            try:
                size = int(kwargs.get('size', 10))
            except Exception:
                raise ValueError('C602')  # Lỗi định dạng lớp trang

            # Parse order
            order_param = kwargs.get('order')
            order = ''
            field_alias = {'id': 'id', 'co': 'code', 'na': 'name', 'de': 'description'}
            if order_param:
                try:
                    order_param = order_param.strip('[]')  # Bỏ []
                    order_parts = order_param.split('-')
                    for part in order_parts:
                        alias, direction = part.split(':') # vdu alias = "co", direction = "1"
                        field = field_alias.get(alias.strip()) # lấy ra value của bí danh
                        if not field or direction not in ['0', '1']:
                            raise Exception()
                        direction = 'desc' if direction == '1' else 'asc'
                        order += f'{field} {direction}, ' # order = 'code desc, name asc, '
                    order = order.strip(', ') #bỏ dấu phẩy ở cuối
                except Exception:
                    raise ValueError('C600')  # Định dạng order sai → lỗi không xác định

            # Parse columnlist
            columnlist = kwargs.get('columnlist')
            if columnlist:
                try:
                    columns = []
                    for alias in columnlist.split(','):
                        field = field_alias.get(alias.strip())
                        if not field:
                            raise Exception()
                        columns.append(field)
                    if 'id' not in columns:
                        columns.insert(0, 'id')
                except Exception:
                    raise ValueError('C607')  # Lỗi columnList
            else:
                columns = ['id', 'code', 'name', 'description']

            # Tìm kiếm
            search = kwargs.get('search')
            domain = []
            if search:
                # danh sách tiêu chí gồm bộ ba  (field_name, operator, value)
                # toán tử logic : prefix form
                domain = ['|', ('name', 'ilike', search), 
               '|', ('description', 'ilike', search), ('code', 'ilike', search)] #ilike : không phân biệt chữ hoa thường

            class_model = request.env['my_api_module.classes']
            # Parse toplist
            toplist_ids = kwargs.get('toplist')
            toplist_records = []
            if toplist_ids:
                try:
                    toplist_ids = [int(x) for x in toplist_ids.split(',') if x.strip().isdigit()]
                    toplist_records = class_model.browse(toplist_ids).read(columns) # trả về list các dict
                except Exception:
                    raise ValueError('C600')  # toplist sai định dạng → lỗi không xác định

            # Phân trang
            total_items = class_model.search_count(domain) # trả về số bản ghi thỏa mãn 
            total_pages = (total_items + size - 1) // size ## làm tròn tính tổng số trang
            offset = (init - 1) * size
            if init > total_pages and total_pages != 0:
                raise ValueError('C601') 
            # offset : số lượng bản ghi bỏ qua, limit : số lượng tối đa bản ghi trả về
            records = class_model.search(domain, order=order or 'id asc', limit=size, offset=offset)
            # Lọc các bản ghi trùng ID với toplist
            toplist_id_set = set(r['id'] for r in toplist_records)
            filtered_records = [r for r in records if r.id not in toplist_id_set]

            # Chuyển sang dict theo các cột yêu cầu
            result_records = [{f: getattr(r, f, '') for f in columns} for r in filtered_records]

            # Gộp lại kết quả cuối cùng
            final_records = toplist_records + result_records

            return request.make_response(
                json.dumps({
                    'code': 200,
                    'status': 'success',
                    'message': 'Lấy danh sách lớp thành công',
                    'data': {
                        'records': final_records,
                        'page_info': {
                            'total_items': total_items,
                            'total_pages': total_pages,
                            'current': init,
                            'size': size
                        }
                    }
                }),
                headers = [('Content-Type', 'application/json')]
            )

        except ValueError as ve:
            error_code = str(ve)
            error_messages = {
                'C601': 'Lỗi định dạng của số trang',
                'C602': 'Lỗi định dạng của Lớp trang',
                'C607': 'Lỗi danh sách columnList'
            }
            return request.make_response(
                json.dumps({
                    'code': error_code,
                    'status': 'error',
                    'message': error_messages.get(error_code, 'Tham số không hợp lệ')
                }),
                headers = [('Content-Type', 'application/json')]
            )

        except Exception as e:
            return http.Response(
                json.dumps({
                    'code': 'C600',
                    'status': 'error',
                    'message': 'Lỗi không xác định: ' + str(e)
                }),
                content_type='application/json'
            )



##3: THÊM MỚI MỘT LỚP

    @http.route('/tra_class', type='http', auth='public', methods=['POST'], csrf=False)
    def store(self, **kw):
        _logger.info(f"Received data: {kw}")
        try:
            # errors = ClassValidator.validate_create(kw, request.env)
            # thêm error validator
            validator = ClassValidator(data = kw,env = request.env['my_api_module.classes'])
            errors = validator.validate_create_data()
            if errors:
                return request.make_response(
                    json.dumps({
                        "code": "E603",
                        "status": "error",
                        "message": "Lỗi kiểm tra dữ liệu : " + " ".join(errors)
                    }),
                    headers=[('Content-Type', 'application/json')]
                )
            # Kiểm tra dữ liệu bắt buộc
            # if not code or not name:
            #     return request.make_response(
            #         json.dumps({
            #             "status": "error",
            #             "code": "E603",
            #             "message": "Lỗi kiểm tra dữ liệu"
            #         }),
            #         headers=[('Content-Type', 'application/json')]
            #     )

            # Kiểm tra trùng mã lớp
            # existing = request.env['my_api_module.classes'].search([('code', '=', code)], limit=1)
            # if existing:
            #     return request.make_response(
            #         json.dumps({
            #             "status": "error",
            #             "code": "E603",
            #             "message": "Lỗi kiểm tra dữ liệu"
            #         }),
            #         headers=[('Content-Type', 'application/json')]
            #     )
            # code = kw.get('code')
            # name = kw.get('name')
            # description = kw.get('description', '')
            # Tạo bản ghi
            # new_class = request.env['my_api_module.classes'].create({
            #     'code': code,
            #     'name': name,
            #     'description': description or ''
            # })
            try:
                new_class = request.env['my_api_module.classes'].create(kw)
            except Exception as e:
                _logger.info(f"Error during create: {str(e)}")
                
            # Trả về kết quả thành công
            return request.make_response(
                json.dumps({
                    "code": 200,
                    "status": "success",
                    "message": "Tạo mới lớp thành công",
                    "data": {
                        "id": new_class.id,
                        "class": new_class.code,
                        "name": new_class.name,
                        "description": new_class.description
                    }
                }),
                headers=[('Content-Type', 'application/json')]
            )
       
        except Exception as e:
            # Không in lỗi chi tiết ra ngoài
            return request.make_response(
                json.dumps({
                    "status": "error",
                    "code": "E600",
                    "message":f"Lỗi không xác định: {str(e)}"
                }),
                headers=[('Content-Type', 'application/json')]
            )

##4: LẤY THÔNG TIN THEO ID
    @http.route('/tra_class/<int:record_id>', type='http', auth='public', methods=['GET'], csrf=False)
    def get_by_id(self, record_id, **kw):
        try:
            # Map bí danh sang field thực tế
            alias_map = {
                'co': 'code',
                'na': 'name',
                'de': 'description',
                'wr': 'write_date'
            }

            # Parse columnlist
            columnlist = kw.get('columnlist')
            if columnlist:
                try:
                    aliases = [alias.strip() for alias in columnlist.split(',')]
                    fields = ['id'] + [alias_map[a] for a in aliases if a in alias_map]
                    if len(fields) == 1:
                        raise ValueError('D607')  # Không có alias hợp lệ
                except Exception:
                    raise ValueError('D607')  # Sai định dạng columnlist
            else:
                fields = ['id', 'code', 'name', 'description']

            # Lấy bản ghi
            class_model = request.env['my_api_module.classes']
            record = class_model.browse(record_id)
            if not record.exists():
                raise ValueError('D604')  # Không tồn tại bản ghi

            # Duyệt lấy dữ liệu các field
            data = {}
            for f in fields:
                value = getattr(record, f, '')
                data[f] = str(value) if value is not None else ''

            # Trả kết quả thành công
            return request.make_response(
                json.dumps({
                    "code": "200",
                    "status": "success",
                    "message": "Lấy thông tin lớp thành công",
                    "data": data
                }),
                headers=[('Content-Type', 'application/json')]
            )

        except ValueError as ve:
            code = str(ve)
            messages = {
                'D604': 'Lỗi: bản ghi không tồn tại',
                'D607': 'Lỗi: danh sách columnList không hợp lệ'
            }
            return request.make_response(
                json.dumps({
                    "code": code,
                    "status": "error",
                    "message": messages.get(code, "Tham số không hợp lệ"),
                    "data": {}
                }),
                headers=[('Content-Type', 'application/json')]
            )

        except Exception as e:
            return request.make_response(
                json.dumps({
                    "code": "D600",
                    "status": "error",
                    "message": "Lỗi không xác định",
                    "data": {}
                }),
                headers=[('Content-Type', 'application/json')]
            )


##5: CẬP NHẬT THÔNG TIN LỚP
    @http.route('/tra_class/<int:id>', type='http', auth='public', methods=['PUT','POST'], csrf=False)
    def update(self, id, **kw):
        _logger.info(f"data nhận được: {kw}")
        #Bản gốc chưa có validator
        # try:
        #     _logger.info(f"data đầu vào: {type(kw)}")
        #     # Tìm bản ghi theo ID
        #     class_record = request.env['my_api_module.classes'].browse(id) # trả về recordset

        #     if not class_record.exists():
        #         return request.make_response(
        #             json.dumps({
        #                 "code": "F603",
        #                 "status": "error",
        #                 "message": "Không tìm thấy bản ghi với ID đã cho",
        #                 "data": {}
        #             }),
        #             headers=[('Content-Type', 'application/json')]
        #         )

        #     # Danh sách các field cho phép cập nhật
        
        #     update_values = {}

        #     for field in allowed_fields:
        #         if field in kw:
        #             update_values[field] = kw.get(field) # lấy giá trị từ request

        #     # Không có dữ liệu cập nhật
        #     if not update_values:
        #         return request.make_response(
        #             json.dumps({
        #                 "code": "F603",
        #                 "status": "error",
        #                 "message": "Không có dữ liệu cần cập nhật",
        #                 "data": {}
        #             }),
        #             headers=[('Content-Type', 'application/json')]
        #         )

        #     # Cập nhật bản ghi
        #     class_record.write(update_values) #update record với tham số là dict

        #     return request.make_response(
        #         json.dumps({
        #             "code": 200,
        #             "status": "success",
        #             "message": "Cập nhật thông tin lớp thành công",
        #             "data": {
        #                 "id": class_record.id,
        #                 "code": class_record.code,
        #                 "name": class_record.name,
        #                 "description": class_record.description
        #             }
        #         }),
        #         headers=[('Content-Type', 'application/json')]
        #     )

        # except Exception as e:
        #     return request.make_response(
        #         json.dumps({
        #             "code": "F600",
        #             "status": "error",
        #             "message": f"Lỗi không xác định: {str(e)}",
        #             "data": {}
        #         }),
        #         headers=[('Content-Type', 'application/json')]
        #     )

        #Bản mới có validator
        allowed_fields = ['code', 'name', 'description']
        try:
        # Tìm bản ghi theo ID
            data = kw
            class_record = request.env['my_api_module.classes'].browse(id)
            if not class_record.exists():
                return request.make_response(
                    json.dumps({
                        "code": "F603",
                        "status": "error",
                        "message": "Không tìm thấy bản ghi với ID đã cho",
                        "data": {}
                    }),
                    headers=[('Content-Type', 'application/json')]
                )

            files = request.httprequest.files.getlist('file')
            file = files[0] if files else None
            timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
            if file:
                filename = f"class_{timestamp}_{class_record.id}_{file.filename}"
                filepath = os.path.join(DOCS_DIR, filename)
                # Xử lý xóa file cũ của id nếu có
                old_files_pattern = os.path.join(DOCS_DIR, f"class_*_{class_record.id}_*")
                old_files = glob.glob(old_files_pattern)
                for old_file in old_files:
                    os.remove(old_file)
                    _logger.info(f"[FILE] Đã xóa file cũ: {old_file}")
                with open(filepath, 'wb') as f:
                    file.save(f)
                _logger.info(f"[FILE] Đã lưu file vào: {filepath}")

                # Đọc file CSV hoặc Excel
                try:
                    if filename.lower().endswith('.csv'):
                        df = pd.read_csv(filepath)
                    elif filename.lower().endswith(('.xls', '.xlsx')):
                        df = pd.read_excel(filepath)
                        if 'dob' in df.columns:
                            df['dob'] = pd.to_datetime(df['dob']).dt.strftime('%Y-%m-%d')
                        _logger.info(f"[DEBUG] DataFrame đọc được:\n{df}")
                    else:
                        raise ValueError("Lỗi kiểm tra dữ liệu")

                    if df.empty:
                        raise ValueError("Lỗi kiểm tra dữ liệu")

                    record = df.iloc[0].to_dict()
                    _logger.info(f"[PARSE FILE] Dữ liệu lấy từ file: {record}")
                    
                    for key, value in record.items():
                        if key not in allowed_fields or pd.isna(value):
                            continue
                        data[key] = value 
                except Exception as e:
                    return request.make_response(json.dumps({
                        "code": "F600",
                        "status": "error",
                        "message": f"Lỗi khi đọc file: {str(e)}",
                        "data": {}
                    }), headers=[('Content-Type', 'application/json')])
            # Gọi validator
            validator = ClassValidator(data, request.env['my_api_module.classes'])
            errors = validator.validate_update_data(id)
            if errors:
                return request.make_response(
                    json.dumps({
                        "code": "F603",
                        "status": "error",
                        "message": "Lỗi kiểm tra dữ liệu: " + " ".join(errors),
                        "data": {}
                    }),
                    headers=[('Content-Type', 'application/json')]
                )

            # Danh sách các field cho phép cập nhật
            # allowed_fields = ['code', 'name', 'description']
            update_values = {}
            for field in allowed_fields:
                if field in data:
                    update_values[field] = data.get(field)

            if not update_values:
                return request.make_response(
                    json.dumps({
                        "code": "F603",
                        "status": "error",
                        "message": "Không có dữ liệu cần cập nhật",
                        "data": {}
                    }),
                    headers=[('Content-Type', 'application/json')]
                )

            # Cập nhật bản ghi
            class_record.write(update_values)

            return request.make_response(
                json.dumps({
                    "code": 200,
                    "status": "success",
                    "message": "Cập nhật thông tin lớp thành công",
                    "data": {
                        "id": class_record.id,
                        "code": class_record.code,
                        "name": class_record.name,
                        "description": class_record.description
                    }
                }),
                headers=[('Content-Type', 'application/json')]
            )

        except Exception as e:
            return request.make_response(
                json.dumps({
                    "code": "F600",
                    "status": "error",
                    "message": f"Lỗi không xác định: {str(e)}",
                    "data": {}
                }),
                headers=[('Content-Type', 'application/json')]
            )
##6: XÓA MỘT LỚP
    @http.route('/tra_class/<int:id>', type='http', auth='public', methods=['DELETE'], csrf=False)
    def destroy(self, id):
        
        try:
            # Tìm bản ghi lớp
            class_record = request.env['my_api_module.classes'].browse(id)
            
            # Kiểm tra tồn tại
            if not class_record.exists():
                return request.make_response(
                    json.dumps({
                        "code": "G604",
                        "status": "error",
                        "message": "Bản ghi không tồn tại",
                        "data": {}
                    }),
                    headers=[('Content-Type', 'application/json')]
                )
            file_pattern = os.path.join(DOCS_DIR, f"class_*_{class_record.id}_*")
            files = glob.glob(file_pattern)
            for file_path in files:
                os.remove(file_path)
                _logger.info(f"[DELETE] Đã xóa file: {file_path}")
            try:
                delete_id = class_record.id
                # Cố gắng xóa bản ghi
                class_record.unlink()

            except ValidationError:
                # Nếu bị khóa ngoại hoặc ràng buộc dữ liệu
                return request.make_response(
                    json.dumps({
                        "code": "G605",
                        "status": "error",
                        "message": "Không thể xóa vì bản ghi có liên kết khóa ngoại",
                        "data": {}
                    }),
                    headers=[('Content-Type', 'application/json')]
                )

            # Trả về kết quả thành công
            return request.make_response(
                json.dumps({
                    "code": 200,
                    "status": "success",
                    "message": "Xóa bản ghi thành công",
                    "data": {"id": delete_id}
                }),
                headers=[('Content-Type', 'application/json')]
            )

        except Exception as e:
            return request.make_response(
                json.dumps({
                    "code": "G600",
                    "status": "error",
                    "message": f"Lỗi không xác định: {str(e)}",
                    "data": {}
                }),
                headers=[('Content-Type', 'application/json')]
            )

##7: SAO CHÉP MỘT LỚP 
    @http.route('/tra_class/copy/<int:id>', type='http', auth='public', methods=['POST'], csrf=False)
    def copy(self, id):
        try:
            # Tìm bản ghi gốc
            original = request.env['my_api_module.classes'].browse(id)

            # Kiểm tra tồn tại
            if not original.exists():
                return request.make_response(
                    json.dumps({
                        "code": "H604",
                        "status": "error",
                        "message": "Bản ghi không tồn tại",
                        "data": {}
                    }),
                    headers=[('Content-Type', 'application/json')]
                )

            # Kiểm tra dữ liệu gốc có hợp lệ không
            if not original.code or not original.name:
                return request.make_response(
                    json.dumps({
                        "code": "H603",
                        "status": "error",
                        "message": "Bản ghi gốc thiếu mã hoặc tên",
                        "data": {}
                    }),
                    headers=[('Content-Type', 'application/json')]
                )

            # Tạo bản ghi mới (sao chép)
            new_class = request.env['my_api_module.classes'].create({
                'code': f'{original.code}',
                'name': f'{original.name}',
                'description': original.description or ''
            })

            new_code = f"{original.code}_{new_class.id}"
            new_class.write({'code': new_code})
            # Trả về kết quả
            return request.make_response(
                json.dumps({
                    "code": 200,
                    "status": "success",
                    "message": "Sao chép bản ghi thành công",
                    "data": {
                        "id": new_class.id,
                        "code": new_class.code,
                        "name": new_class.name,
                        "description": new_class.description
                    }
                }),
                headers=[('Content-Type', 'application/json')]
            )

        except Exception as e:
            return request.make_response(
                json.dumps({
                    "code": "H600",
                    "status": "error",
                    "message": f"Lỗi không xác định: {str(e)}",
                    "data": {}
                }),
                headers=[('Content-Type', 'application/json')]
            )

##8: XÓA NHIỀU LỚP
    @http.route('/tra_class/delete', type='http', auth='public', methods=['DELETE'], csrf=False)
    def mass_delete(self, **kw):
        try:
            idlist_raw = kw.get('idlist')
            if not idlist_raw:
                raise ValueError('I604')  # Thiếu tham số idlist

            idlist = [int(i.strip()) for i in idlist_raw.split(',') if i.strip().isdigit()]
            if not idlist:
                raise ValueError('I604')  # Danh sách rỗng hoặc sai định dạng

            class_model = request.env['my_api_module.classes']
            records = class_model.browse(idlist).exists()  # <-- Chỉ giữ bản ghi tồn tại thật sự

            existing_ids = records.mapped('id')
            not_found_ids = [i for i in idlist if i not in existing_ids]
            if not_found_ids:
                raise ValueError(f"I604 - Không tìm thấy các ID: {not_found_ids}")
            #Xóa bản ghi trước
            for class_record in records:
                file_pattern = os.path.join(DOCS_DIR, f"class_*_{class_record.id}_*")
                files = glob.glob(file_pattern)
                for file_path in files:
                    os.remove(file_path)
                    _logger.info(f"[DELETE] Đã xóa file: {file_path}")
            try:
                records.unlink()
            except Exception:
                raise ValueError('I604')  # Có thể do ràng buộc khóa ngoại

            return request.make_response(
                json.dumps({
                    "code": 200,
                    "status": "success",
                    "message": "Xóa các bản ghi thành công",
                    "data": {"ids": existing_ids}
                }),
                headers=[('Content-Type', 'application/json')]
            )

        except ValueError as ve:
            return request.make_response(
                json.dumps({
                    "code": str(ve),
                    "status": "error",
                    "message": "Lỗi các bản ghi không xóa được.",
                    "data": {}
                }),
                headers=[('Content-Type', 'application/json')]
            )

        except Exception as e:
            return request.make_response(
                json.dumps({
                    "code": "I600",
                    "status": "error",
                    "message": f"Lỗi không xác định: {str(e)}",
                    "data": {}
                }),
                headers=[('Content-Type', 'application/json')]
            )


##9: COPY NHIỀU LỚP
    @http.route('/tra_class/copy', type='http', auth='public', methods=['POST'], csrf=False)
    def mass_copy(self, **kw):
        try:
            idlist_str = kw.get('idlist')
            if not idlist_str:
                raise ValueError('H603')  # Thiếu dữ liệu

            # Parse idlist
            idlist = [int(x) for x in idlist_str.split(',') if x.strip().isdigit()]
            if not idlist:
                raise ValueError('H603')  # Dữ liệu không hợp lệ

            class_model = request.env['my_api_module.classes']
            records = class_model.browse(idlist)

            # Kiểm tra xem có bản ghi nào không tồn tại
            existing_ids = records.ids
            missing_ids = [rid for rid in idlist if rid not in existing_ids]
            if missing_ids:
                raise ValueError('H604')  # Lỗi kiểm tra dữ liệu

            copied_data = []
            for record in records:
                # Bước 1: Tạo bản ghi với code tạm thời
                temp_code = record.code or 'copy'
                new_record = class_model.create({
                    'code': temp_code,
                    'name': record.name,
                    'description': record.description or '',
                })

                # Bước 2: Cập nhật lại code theo format: <code gốc>_<id mới>
                final_code = f"{record.code}_{new_record.id}" if record.code else f"copy_{new_record.id}"
                new_record.write({'code': final_code})

                copied_data.append({
                    'id': new_record.id,
                    'code': new_record.code,
                    'name': new_record.name,
                    'description': new_record.description
                })

            return request.make_response(
                json.dumps({
                    "code": 200,
                    "status": "success",
                    "message": "Sao chép thành công",
                    "data": copied_data
                }),
                headers=[('Content-Type', 'application/json')]
            )

        except ValueError as ve:
            code = str(ve)
            messages = {
                'H603': "Lỗi kiểm tra dữ liệu: idlist không hợp lệ hoặc bản ghi không tồn tại",
                'H604': "Lỗi kiểm tra dữ liệu: không tìm thấy bản ghi với ID đã cho"
            }
            return request.make_response(
                json.dumps({
                    "code": code,
                    "status": "error",
                    "message": messages.get(code, "Tham số không hợp lệ"),
                    "data": []
                }),
                headers=[('Content-Type', 'application/json')]
            )
        except Exception:
            return request.make_response(
                json.dumps({
                    "code": "H600",
                    "status": "error",
                    "message": "Lỗi không xác định",
                    "data": []
                }),
                headers=[('Content-Type', 'application/json')]
            )

##10: Import bản ghi
    @http.route('/tra_class/import', type='http', auth='public', methods=['POST'], csrf=False)
    def import_data(self, **kw):
        try:
            # ==== Kiểm tra file đính kèm ====
            file = kw.get('attachment')
            if not file:
                raise ValueError("J604")

            filename = file.filename.lower()
            if not (filename.endswith('.csv') or filename.endswith('.xlsx') or filename.endswith('.xls')):
                raise ValueError("J601")

            # ==== Đọc dữ liệu từ file ====
            records = []
            if filename.endswith('.csv'):
                content = file.stream.read().decode('utf-8')
                reader = csv.DictReader(io.StringIO(content)) ## trả về dictionary với key là tên cột
                for row in reader:
                    records.append(row)

            elif filename.endswith(('.xlsx', '.xls')):
                if not pd:
                    raise ImportError("Chưa cài pandas hoặc openpyxl.")
                df = pd.read_excel(file.stream) #đọc file dataframe
                records = df.to_dict(orient='records')

            if not records:
                raise ValueError("J605")

            # ==== Thêm bản ghi mới vào hệ thống ====
            created = []
                        # ==== Thêm bản ghi mới vào hệ thống ==== 
            created = []
            for row in records:
                code = row.get('code')
                name = row.get('name')
                description = row.get('description', '')

                if not code or not name:
                    continue  # Bỏ qua bản ghi không đủ thông tin

                # Kiểm tra xem code đã tồn tại chưa
                existing = request.env['my_api_module.classes'].search([('code', '=', code)], limit=1)

                # Nếu trùng thì để tạm code rỗng, lát nữa gán sau khi có ID
                temp_code = code if not existing else ''

                # Tạo bản ghi
                new_class = request.env['my_api_module.classes'].create({
                    'code': temp_code,
                    'name': name,
                    'description': description,
                })

                # Nếu bị trùng, cập nhật lại code với _id
                if not temp_code:
                    new_code = f"{code}_{new_class.id}"
                    new_class.write({'code': new_code})
                else:
                    new_code = code

                created.append({
                    'id': new_class.id,
                    'code': new_code,
                    'name': new_class.name,
                    'description': new_class.description
                })


            if not created:
                raise ValueError("J605")

            # ==== Trả về dữ liệu thành công ====
            return request.make_response(
                json.dumps({
                    "code": "200",
                    "status": "success",
                    "message": "Import thành công",
                    "data": created
                }, ensure_ascii=False),
                headers=[('Content-Type', 'application/json')]
            )

        except ValueError as ve:
            code = str(ve)
            messages = {
                'J601': "Lỗi định dạng tệp. Chỉ hỗ trợ CSV hoặc Excel.",
                'J604': "Lỗi tệp dữ liệu chưa tải lên.",
                'J605': "Lỗi không có dữ liệu được thêm mới từ file."
            }
            return request.make_response(
                json.dumps({
                    "code": code,
                    "status": "error",
                    "message": messages.get(code, f"Lỗi: {code}"),
                    "data": []
                }, ensure_ascii=False),
                headers=[('Content-Type', 'application/json')]
            )

        except Exception as e:
            return request.make_response(
                json.dumps({
                    "code": "J600",
                    "status": "error",
                    "message": f"Lỗi không xác định: {str(e)}",
                    "data": []
                }, ensure_ascii=False),
                headers=[('Content-Type', 'application/json')]
            )


##11: Xuất một bản ghi

    @http.route('/tra_class/export/<int:record_id>', type='http', auth='public', methods=['GET'], csrf=False)
    def export_by_id(self, record_id, **kw):
        try:
            export_type = kw.get('type', 'csv').lower()
            columnlist = kw.get('columnlist')

            if not columnlist:
                raise ValueError("L607")

            # Danh sách bí danh người dùng truyền vào
            aliases = [a.strip() for a in columnlist.split(',')]

            # Mapping bí danh → field thật trong model
            alias_to_field = {
                'co': 'code',
                'na': 'name',
                'de': 'description',
                'wr': 'write_date',
                # Thêm các alias khác tại đây nếu cần
            }

            # Duyệt và ánh xạ sang field thực
            fields = []
            for alias in aliases:
                if alias not in alias_to_field:
                    raise ValueError(f"Bí danh '{alias}' không hợp lệ.")
                fields.append(alias_to_field[alias])

            if export_type not in ['csv', 'xlsx', 'docx']:
                raise ValueError("L601")

            class_model = request.env['my_api_module.classes']
            
                
            record = class_model.browse(record_id)

            if not record.exists():
                raise ValueError("L602")

            # Lấy dữ liệu theo các field thực
            row = []
            for field in fields:
                value = getattr(record, field, '')
                if hasattr(value, 'name'):
                    value = value.name
                elif hasattr(value, 'isoformat'):
                    value = value.isoformat()
                elif isinstance(value, (list, tuple)):
                    value = ', '.join(str(v) for v in value)
                elif isinstance(value, bool):
                    value = 'Có' if value else 'Không'
                elif value is False or value is None:
                    value = ''
                row.append(value)

            # ===== CSV Export =====
            if export_type == 'csv':
                buffer = io.StringIO()
                writer = csv.writer(buffer)
                writer.writerow(fields)  # Dùng tên field thật làm header
                writer.writerow(row)
                content = buffer.getvalue()
                buffer.close()
                response_data = content.encode('utf-8')
                content_type = 'text/csv'

            # ===== XLSX Export =====
            elif export_type == 'xlsx':
                if not xlsxwriter:
                    raise ImportError("Chưa cài thư viện xlsxwriter.")
                output = io.BytesIO()
                workbook = xlsxwriter.Workbook(output, {'in_memory': True})
                worksheet = workbook.add_worksheet()
                for col_num, header in enumerate(fields):
                    worksheet.write(0, col_num, header)
                for col_num, cell in enumerate(row):
                    worksheet.write(1, col_num, cell)
                workbook.close()
                response_data = output.getvalue()
                output.close()
                content_type = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'

            # ===== DOCX Export =====
            elif export_type == 'docx':
                try:
                    from docx import Document
                except ImportError:
                    raise ImportError("Chưa cài thư viện python-docx.")
                doc = Document()
                doc.add_heading('Thông tin lớp', 0)
                table = doc.add_table(rows=1, cols=len(fields))
                hdr_cells = table.rows[0].cells
                for idx, field in enumerate(fields):
                    hdr_cells[idx].text = field  # Dùng tên field thật
                row_cells = table.add_row().cells
                for idx, cell in enumerate(row):
                    row_cells[idx].text = str(cell)
                output = io.BytesIO()
                doc.save(output)
                response_data = output.getvalue()
                output.close()
                content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

            return request.make_response(
                response_data,
                headers=[
                    ('Content-Type', content_type),
                    ('Content-Disposition', f'attachment; filename=class_{record_id}.{export_type}')
                ]
            )

        except ValueError as ve:
            code = str(ve)
            messages = {
                'L601': "Lỗi định dạng tệp. Chỉ hỗ trợ CSV, XLSX hoặc DOCX.",
                'L602': "Không tìm thấy bản ghi lớp cần xuất.",
                'L607': "Danh sách columnlist không hợp lệ."
            }
            return request.make_response(
                json.dumps({
                    "code": code,
                    "status": "error",
                    "message": messages.get(code, f"Lỗi: {code}"),
                }),
                headers=[('Content-Type', 'application/json')]
            )
        except Exception as e:
            return request.make_response(
                json.dumps({
                    "code": "L600",
                    "status": "error",
                    "message": f"Lỗi không xác định: {str(e)}"
                }),
                headers=[('Content-Type', 'application/json')]
            )



####12: Xuất nhiều bản ghi

    @http.route('/tra_class/export', type='http', auth='public', methods=['POST'], csrf=False)
    def mass_export(self, **kw):
        try:
            idlist = kw.get('idlist')
            export_type = kw.get('type', 'csv').lower()
            columnlist = kw.get('columnlist')

            if not idlist:
                raise ValueError("L602")

            if not columnlist:
                raise ValueError("L607")

            # Danh sách bí danh → field thật
            alias_to_field = {
                'co': 'code',
                'na': 'name',
                'de': 'description',
                'wr': 'write_date',
            }

            aliases = [a.strip() for a in columnlist.split(',')]
            fields = []
            for alias in aliases:
                if alias not in alias_to_field:
                    raise ValueError(f"Bí danh '{alias}' không hợp lệ.")
                fields.append(alias_to_field[alias])

            if export_type not in ['csv', 'xlsx', 'docx']:
                raise ValueError("L601")

            ids = [int(i.strip()) for i in idlist.split(',') if i.strip().isdigit()]
            records = request.env['my_api_module.classes'].browse(ids)
            missing_ids = [str(rid) for rid, rec in zip(ids, records) if not rec.exists()]
            if missing_ids:
                raise ValueError(f"L602: ID không tồn tại: {', '.join(missing_ids)}")


            if not records:
                raise ValueError("L602")

            # ===== XỬ LÝ DỮ LIỆU =====
            data_rows = []
            for record in records:
                row = []
                for field in fields:
                    value = getattr(record, field, '')
                    if hasattr(value, 'name'):
                        value = value.name
                    elif hasattr(value, 'isoformat'):
                        value = value.isoformat()
                    elif isinstance(value, (list, tuple)):
                        value = ', '.join(str(v) for v in value)
                    elif isinstance(value, bool):
                        value = 'Có' if value else 'Không'
                    elif value is False or value is None:
                        value = ''
                    row.append(value)
                data_rows.append(row)

            # ===== EXPORT =====
            if export_type == 'csv':
                buffer = io.StringIO()
                writer = csv.writer(buffer)
                writer.writerow(fields)  # Header
                writer.writerows(data_rows)
                content = buffer.getvalue()
                buffer.close()
                response_data = content.encode('utf-8')
                content_type = 'text/csv'

            elif export_type == 'xlsx':
                if not xlsxwriter:
                    raise ImportError("Chưa cài thư viện xlsxwriter.")
                output = io.BytesIO()
                workbook = xlsxwriter.Workbook(output, {'in_memory': True})
                worksheet = workbook.add_worksheet()
                for col_num, header in enumerate(fields):
                    worksheet.write(0, col_num, header)
                for row_idx, row in enumerate(data_rows, start=1):
                    for col_idx, cell in enumerate(row):
                        worksheet.write(row_idx, col_idx, cell)
                workbook.close()
                response_data = output.getvalue()
                output.close()
                content_type = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'

            elif export_type == 'docx':
                try:
                    from docx import Document
                except ImportError:
                    raise ImportError("Chưa cài thư viện python-docx.")
                doc = Document()
                doc.add_heading('Danh sách lớp học', 0)
                table = doc.add_table(rows=1, cols=len(fields))
                hdr_cells = table.rows[0].cells
                for idx, field in enumerate(fields):
                    hdr_cells[idx].text = field
                for row in data_rows:
                    row_cells = table.add_row().cells
                    for idx, cell in enumerate(row):
                        row_cells[idx].text = str(cell)
                output = io.BytesIO()
                doc.save(output)
                response_data = output.getvalue()
                output.close()
                content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

            # ===== TRẢ KẾT QUẢ =====
            return request.make_response(
                response_data,
                headers=[
                    ('Content-Type', content_type),
                    ('Content-Disposition', f'attachment; filename=classes_export.{export_type}')
                ]
            )

        except ValueError as ve:
            code = str(ve)
            messages = {
                'L601': "Lỗi định dạng tệp. Chỉ hỗ trợ CSV, XLSX hoặc DOCX.",
                'L602': "Lỗi danh sách bản ghi không tồn tại.",
                'L607': "Lỗi danh sách columnList.",
            }
            return request.make_response(
                json.dumps({
                    "code": code,
                    "status": "error",
                    "message": messages.get(code, f"Lỗi: {code}")
                }),
                headers=[('Content-Type', 'application/json')]
            )

        except Exception as e:
            return request.make_response(
                json.dumps({
                    "code": "L600",
                    "status": "error",
                    "message": f"Lỗi không xác định: {str(e)}"
                }),
                headers=[('Content-Type', 'application/json')]
            )

